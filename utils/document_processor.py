import io
import re
import os
import tempfile
import traceback
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table
import unicodedata

def enhanced_chunk_docx(file_content, chunk_size=800):
    """Enhanced DOCX chunker with complete content extraction"""
    try:
        # Create document from uploaded content
        doc = Document(io.BytesIO(file_content))
        
        chunks = []
        current_chunk = ""
        chunk_id = 0
        image_counter = 1
        current_context = {'state': None, 'section': None, 'topic': None}
        
        # Create temp directory for images
        temp_image_dir = tempfile.mkdtemp()
        
        def update_context(text):
            """Update current context based on text content"""
            text_upper = text.upper()
            
            # State detection
            state_patterns = {
                'OH': [r'\bOHIO\b', r'\bOH\b(?!\w)'],
                'MD': [r'\bMARYLAND\b', r'\bMD\b(?!\w)'],
                'NJ': [r'\bNEW\s+JERSEY\b', r'\bNJ\b(?!\w)'],
                'IL': [r'\bILLINOIS\b', r'\bIL\b(?!\w)'],
                'NY': [r'\bNEW\s+YORK\b', r'\bNY\b(?!\w)'],
                'NV': [r'\bNEVADA\b', r'\bNV\b(?!\w)'],
                'MA': [r'\bMASSACHUSETTS\b', r'\bMA\b(?!\w)']
            }
            
            for state, patterns in state_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, text_upper):
                        current_context['state'] = state
                        if 'RISE' in text_upper:
                            current_context['section'] = 'RISE'
                        elif 'REGULAR' in text_upper:
                            current_context['section'] = 'REGULAR'
                        break
            
            # Topic detection
            if 'PRICING' in text_upper or 'MENU PRICE' in text_upper:
                current_context['topic'] = 'PRICING'
            elif 'BATTER' in text_upper:
                current_context['topic'] = 'BATTERIES'
            elif 'BATCH SUB' in text_upper:
                current_context['topic'] = 'BATCH_SUB'
            elif 'DELIVERY DATE' in text_upper:
                current_context['topic'] = 'DELIVERY_DATE'
            elif 'ORDER LIMIT' in text_upper:
                current_context['topic'] = 'ORDER_LIMIT'
        
        def clean_caption(text):
            """Enhanced text cleaning with better normalization"""
            cleaned = unicodedata.normalize('NFKC', text)
            cleaned = re.sub(r"\s+", " ", cleaned).strip()
            cleaned = cleaned.replace("–", "-").replace("—", "-").replace(""", '"').replace(""", '"')
            cleaned = cleaned.replace("'", "'").replace("'", "'")
            # Remove excessive punctuation
            cleaned = re.sub(r'[.]{2,}', '.', cleaned)
            return cleaned

        def extract_label(text):
            """Enhanced label extraction supporting multiple formats"""
            text = clean_caption(text)
            
            # Pattern for "Image X: description" format
            caption_pattern = re.compile(r"^Image\s+(\d+)\s*[:.]?\s*(.*?)(?:\.|$)", re.IGNORECASE)
            figure_pattern = re.compile(r"^Figure\s+(\d+)\s*[:.]?\s*(.*?)(?:\.|$)", re.IGNORECASE)
            
            # Try Image pattern first
            m = caption_pattern.match(text)
            if m:
                idx = int(m.group(1))
                desc = m.group(2).strip().rstrip(".")
                return f"Image {idx}: {desc}" if desc else f"Image {idx}"
            
            # Try Figure pattern
            m = figure_pattern.match(text)
            if m:
                idx = int(m.group(1))
                desc = m.group(2).strip().rstrip(".")
                return f"Figure {idx}: {desc}" if desc else f"Figure {idx}"
            
            # Look for descriptive patterns without numbers
            descriptive_patterns = [
                r'([^.]+\s+example\s*[^.]*)',
                r'([^.]+\s+sheet\s*[^.]*)',
                r'([^.]+\s+form\s*[^.]*)',
                r'([^.]+\s+format\s*[^.]*)',
                r'([^.]+\s+setup\s*[^.]*)',
                r'([^.]+\s+process\s*[^.]*)',
                r'([^.]+\s+workflow\s*[^.]*)',
                r'([^.]+\s+template\s*[^.]*)'
            ]
            
            for pattern in descriptive_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    desc = match.group(1).strip()
                    if len(desc) > 5 and len(desc) < 80:  # Reasonable caption length
                        return desc
            
            return None

        def extract_images_with_enhanced_labels():
            """Extract images with sophisticated caption matching"""
            nonlocal image_counter
            
            # Collect all document items in order
            items = []
            body = doc.element.body
            position = 0
            
            for child in body.iterchildren():
                if isinstance(child, CT_P):
                    para = Paragraph(child, doc)
                    
                    # Check for images in paragraph
                    has_image = False
                    for run in para.runs:
                        if 'graphic' in run._element.xml:
                            for drawing in run._element.findall(".//w:drawing", namespaces=run._element.nsmap):
                                for blip in drawing.findall(".//a:blip", namespaces=run._element.nsmap):
                                    rel_id = blip.get(qn('r:embed'))
                                    if rel_id and rel_id in doc.part.related_parts:
                                        image_part = doc.part.related_parts[rel_id]
                                        items.append({
                                            'type': 'image', 
                                            'content': image_part, 
                                            'position': position,
                                            'paragraph_text': para.text.strip(),
                                            'element': child
                                        })
                                        has_image = True
                    
                    # Add text if it exists and doesn't have images
                    if para.text.strip() and not has_image:
                        items.append({
                            'type': 'text', 
                            'content': para.text.strip(), 
                            'position': position
                        })
                        
                elif isinstance(child, CT_Tbl):
                    table = Table(child, doc)
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                # Check for images in table cells
                                for run in para.runs:
                                    if 'graphic' in run._element.xml:
                                        for drawing in run._element.findall(".//w:drawing", namespaces=run._element.nsmap):
                                            for blip in drawing.findall(".//a:blip", namespaces=run._element.nsmap):
                                                rel_id = blip.get(qn('r:embed'))
                                                if rel_id and rel_id in doc.part.related_parts:
                                                    image_part = doc.part.related_parts[rel_id]
                                                    items.append({
                                                        'type': 'image', 
                                                        'content': image_part, 
                                                        'position': position,
                                                        'paragraph_text': para.text.strip(),
                                                        'element': para._element
                                                    })
                
                position += 1

            # Process images with enhanced caption matching
            images = []
            i = 0
            
            while i < len(items):
                if items[i]['type'] == 'image':
                    image_part = items[i]['content']
                    
                    # Look for caption in multiple places
                    label = None
                    
                    # 1. Check if the image's paragraph contains caption text
                    if items[i].get('paragraph_text'):
                        potential_label = extract_label(items[i]['paragraph_text'])
                        if potential_label:
                            label = potential_label
                    
                    # 2. Look ahead for following caption
                    if not label:
                        for j in range(i + 1, min(i + 4, len(items))):  # Look ahead up to 3 items
                            if items[j]['type'] == 'text':
                                potential_label = extract_label(items[j]['content'])
                                if potential_label:
                                    label = potential_label
                                    break
                    
                    # 3. Look behind for preceding caption
                    if not label:
                        for j in range(max(0, i - 3), i):  # Look behind up to 3 items
                            if items[j]['type'] == 'text':
                                potential_label = extract_label(items[j]['content'])
                                if potential_label:
                                    label = potential_label
                                    break
                    
                    # Default label if none found
                    if not label:
                        label = f"Image {image_counter}"
                    
                    # Save image file
                    image_extension = image_part.content_type.split('/')[-1]
                    if image_extension == 'jpeg':
                        image_extension = 'jpg'
                    elif image_extension not in ['jpg', 'png', 'gif', 'bmp', 'webp']:
                        image_extension = 'png'  # Default fallback
                        
                    image_filename = f"image_{image_counter}.{image_extension}"
                    image_path = os.path.join(temp_image_dir, image_filename)
                    
                    with open(image_path, "wb") as f:
                        f.write(image_part.blob)
                    
                    images.append({
                        'filename': image_filename,
                        'path': image_path,
                        'label': label,
                        'number': image_counter,
                        'position': items[i]['position']
                    })
                    
                    image_counter += 1
                    
                i += 1
            
            return images
        
        def process_table(table):
            """Extract table content"""
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(' | '.join(row_text))
            return '\n'.join(table_text)
        
        # Extract all images with enhanced labels first
        all_images = extract_images_with_enhanced_labels()
        
        # Process document body elements in order
        all_content = []
        position = 0
        
        # Process paragraphs and tables
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # Paragraph
                para = next(p for p in doc.paragraphs if p._element == element)
                text = para.text.strip()
                if text:
                    all_content.append({
                        'type': 'paragraph',
                        'text': text,
                        'position': position
                    })
            elif element.tag.endswith('tbl'):
                # Table
                table = next(t for t in doc.tables if t._element == element)
                table_text = process_table(table)
                if table_text:
                    all_content.append({
                        'type': 'table',
                        'text': table_text,
                        'position': position
                    })
            position += 1
        
        # Also check headers and footers
        for section in doc.sections:
            # Headers
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_content.insert(0, {
                            'type': 'header',
                            'text': text,
                            'position': -1  # Headers come first
                        })
            
            # Footers
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_content.append({
                            'type': 'footer',
                            'text': text,
                            'position': 999  # Footers come last
                        })
        
        # Create image lookup by position for efficient matching
        image_by_position = {}
        for img in all_images:
            pos = img.get('position', 0)
            if pos not in image_by_position:
                image_by_position[pos] = []
            image_by_position[pos].append(img)
        
        # Create chunks with strict position-based image attachment
        chunk_ranges = []  # Track [start_pos, end_pos] for each chunk
        current_position = 0
        
        for content in all_content:
            text = content['text']
            content_position = content.get('position', current_position)
            
            if not text:
                continue
            
            # Update context
            update_context(text)
            
            # If adding this text would exceed chunk size, save current chunk
            if len(current_chunk) + len(text) > chunk_size and current_chunk:
                # Record the position range for this chunk
                chunk_ranges.append([current_position - 10, current_position])  # Rough range
                
                chunks.append({
                    'chunk_id': chunk_id,
                    'text': current_chunk.strip(),
                    'images': [],  # Will be filled in next step
                    'start_pos': current_position - 10,
                    'end_pos': current_position,
                    'metadata': {
                        'states': [current_context['state']] if current_context['state'] else [],
                        'sections': [current_context['section']] if current_context['section'] else [],
                        'topics': [current_context['topic']] if current_context['topic'] else [],
                        'word_count': len(current_chunk.split()),
                        'has_images': False,
                        'image_count': 0
                    }
                })
                chunk_id += 1
                current_chunk = ""
            
            current_chunk += text + " "
            current_position = content_position
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                'chunk_id': chunk_id,
                'text': current_chunk.strip(),
                'images': [],
                'start_pos': current_position - 10,
                'end_pos': current_position + 10,
                'metadata': {
                    'states': [current_context['state']] if current_context['state'] else [],
                    'sections': [current_context['section']] if current_context['section'] else [],
                    'topics': [current_context['topic']] if current_context['topic'] else [],
                    'word_count': len(current_chunk.split()),
                    'has_images': False,
                    'image_count': 0
                }
            })
        
        # Now assign images to chunks based on strict document position adjacency
        for chunk in chunks:
            chunk_images = []
            start_pos = chunk.get('start_pos', 0)
            end_pos = chunk.get('end_pos', 999)
            
            # Find images that are immediately before, within, or after this chunk
            # Check positions: [start-3, start-2, start-1, start...end, end+1, end+2, end+3]
            for check_pos in range(start_pos - 3, end_pos + 4):
                if check_pos in image_by_position:
                    chunk_images.extend(image_by_position[check_pos])
            
            # Sort images by their document position to maintain order
            chunk_images.sort(key=lambda x: x.get('position', 0))
            
            # Update chunk with images
            chunk['images'] = chunk_images
            chunk['metadata']['has_images'] = len(chunk_images) > 0
            chunk['metadata']['image_count'] = len(chunk_images)
            
            # Clean up temporary position fields
            if 'start_pos' in chunk:
                del chunk['start_pos']
            if 'end_pos' in chunk:
                del chunk['end_pos']
        
        return chunks
        
    except Exception as e:
        print(f"Error processing document: {e}")
        print(traceback.format_exc())
        return []